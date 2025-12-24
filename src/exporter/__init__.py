import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import List
from pathlib import Path
from loguru import logger
import config

class ExcelExporter:
    """Xuất dữ liệu ra file Excel"""
    
    @staticmethod
    def export_invoices(invoices: List, output_path: str = None) -> str:
        """
        Xuất danh sách hóa đơn ra Excel
        
        Args:
            invoices: Danh sách invoice objects
            output_path: Đường dẫn file output (optional)
            
        Returns:
            Đường dẫn file Excel đã tạo
        """
        try:
            # Convert invoices to dataframe
            data = []
            for inv in invoices:
                data.append({
                    'Số hóa đơn': inv.invoice_number,
                    'Ngày': inv.invoice_date.strftime('%d/%m/%Y') if inv.invoice_date else '',
                    'Nhà cung cấp': inv.supplier_name,
                    'MST': inv.supplier_tax_code,
                    'Địa chỉ': inv.supplier_address,
                    'Tiền trước thuế': inv.subtotal,
                    'Thuế suất (%)': inv.tax_rate,
                    'Tiền thuế': inv.tax_amount,
                    'Tổng tiền': inv.total_amount,
                    'Mô tả': inv.description,
                    'Tài khoản': inv.account_code,
                    'Danh mục': inv.category,
                    'Trạng thái': inv.status,
                    'Người tạo': inv.created_by_username,
                    'Ngày tạo': inv.created_at.strftime('%d/%m/%Y %H:%M') if inv.created_at else ''
                })
            
            df = pd.DataFrame(data)
            
            # Generate output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = config.DATA_DIR / f'invoices_export_{timestamp}.xlsx'
            
            # Create Excel file with formatting
            wb = Workbook()
            ws = wb.active
            ws.title = "Hóa đơn"
            
            # Write headers
            headers = list(df.columns)
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_num, value=header)
                cell.font = Font(bold=True, color="FFFFFF", size=11)
                cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )
            
            # Write data
            for row_num, row_data in enumerate(df.values, 2):
                for col_num, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row_num, column=col_num, value=value)
                    cell.border = Border(
                        left=Side(style='thin'),
                        right=Side(style='thin'),
                        top=Side(style='thin'),
                        bottom=Side(style='thin')
                    )
                    
                    # Format currency columns
                    if col_num in [6, 8, 9]:  # Subtotal, Tax amount, Total
                        cell.number_format = '#,##0'
            
            # Adjust column widths
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column].width = adjusted_width
            
            # Add summary row
            summary_row = len(df) + 3
            ws.cell(row=summary_row, column=1, value="TỔNG CỘNG").font = Font(bold=True)
            ws.cell(row=summary_row, column=6, value=df['Tiền trước thuế'].sum()).font = Font(bold=True)
            ws.cell(row=summary_row, column=8, value=df['Tiền thuế'].sum()).font = Font(bold=True)
            ws.cell(row=summary_row, column=9, value=df['Tổng tiền'].sum()).font = Font(bold=True)
            
            # Save file
            wb.save(output_path)
            logger.info(f"Exported {len(invoices)} invoices to Excel: {output_path}")
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error exporting to Excel: {e}")
            raise

class WordExporter:
    """Xuất dữ liệu ra file Word"""
    
    @staticmethod
    def export_invoice_report(invoices: List, output_path: str = None) -> str:
        """
        Xuất báo cáo hóa đơn ra Word
        
        Args:
            invoices: Danh sách invoice objects
            output_path: Đường dẫn file output (optional)
            
        Returns:
            Đường dẫn file Word đã tạo
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('BÁO CÁO HÓA ĐƠN', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date range
            date_range = doc.add_paragraph()
            date_range.add_run(f'Ngày xuất báo cáo: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
            date_range.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Blank line
            
            # Summary statistics
            doc.add_heading('Tổng quan', 1)
            total_invoices = len(invoices)
            total_amount = sum(inv.total_amount for inv in invoices)
            
            summary = doc.add_paragraph()
            summary.add_run(f'Tổng số hóa đơn: ').bold = True
            summary.add_run(f'{total_invoices}\n')
            summary.add_run(f'Tổng giá trị: ').bold = True
            summary.add_run(f'{total_amount:,.0f} VNĐ\n')
            
            doc.add_paragraph()  # Blank line
            
            # Invoice details
            doc.add_heading('Chi tiết hóa đơn', 1)
            
            # Create table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Số HĐ', 'Ngày', 'Nhà cung cấp', 'Mô tả', 'Tài khoản', 'Tổng tiền']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for inv in invoices:
                row_cells = table.add_row().cells
                row_cells[0].text = inv.invoice_number
                row_cells[1].text = inv.invoice_date.strftime('%d/%m/%Y') if inv.invoice_date else ''
                row_cells[2].text = inv.supplier_name
                row_cells[3].text = inv.description[:50] + '...' if len(inv.description) > 50 else inv.description
                row_cells[4].text = inv.account_code or ''
                row_cells[5].text = f'{inv.total_amount:,.0f}'
            
            # Generate output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_path = config.DATA_DIR / f'invoices_report_{timestamp}.docx'
            
            # Save file
            doc.save(output_path)
            logger.info(f"Exported invoice report to Word: {output_path}")
            
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error exporting to Word: {e}")
            raise

class StatisticsExporter:
    """Xuất thống kê và báo cáo"""
    
    @staticmethod
    def generate_monthly_summary(invoices: List) -> dict:
        """Tạo báo cáo tổng hợp theo tháng"""
        df = pd.DataFrame([inv.to_dict() for inv in invoices])
        
        summary = {
            'total_invoices': len(invoices),
            'total_amount': df['total_amount'].sum(),
            'average_amount': df['total_amount'].mean(),
            'by_category': df.groupby('category')['total_amount'].sum().to_dict(),
            'by_account': df.groupby('account_code')['total_amount'].sum().to_dict(),
            'by_supplier': df.groupby('supplier_name')['total_amount'].sum().to_dict()
        }
        
        return summary
